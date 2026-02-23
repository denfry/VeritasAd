"""
Улучшенный скрипт для заполнения документов по преддипломной практике
Проект: VeritasAd - AI платформа детекции скрытой рекламы
Студент: Юрков Данила Андреевич, группа МВА-122
Период практики: 09.02.2026 - 22.02.2026
"""

import os
import sys
# Устанавливаем UTF-8 кодировку для вывода в Windows
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

# Константы
STUDENT_FIO = "Юрков Данила Андреевич"
STUDENT_GROUP = "МВА-122"
THEMA_PRACTICE = "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
THEMA_VKR = "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
ORGANIZATION = "VeritasAd"
START_DATE = datetime(2026, 2, 9)
END_DATE = datetime(2026, 2, 22)
SUPERVISOR = "Беспалов М.Е., к.т.н., доц."
HEAD_DEPARTMENT = "Травкин Е.И., к.п.н."

DOCS_DIR = r"C:\Users\dabin\Documents\Projects\VeritasAd\docs\диплом практика"


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Замена текста в параграфе с сохранением форматирования"""
    if old_text not in paragraph.text:
        return False
    
    # Создаем новую строку с заменой
    new_paragraph_text = paragraph.text.replace(old_text, new_text)
    
    # Очищаем все runs кроме первого
    runs = paragraph.runs
    if runs:
        # Сохраняем форматирование первого run
        first_run = runs[0]
        for run in runs[1:]:
            run.text = ""
        runs[0].text = new_paragraph_text
    
    return True


def fill_title_page():
    """Заполнение титульного листа отчета (Форма А)"""
    print("Заполнение титульного листа отчета...")
    
    doc_path = os.path.join(DOCS_DIR, "4-Титульный лист отчета преддипломной практики Форма А.docx")
    
    try:
        doc = Document(doc_path)
        
        # Проходим по всем параграфам и ищем ключевые фразы для замены
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Заполняем ФИО
            if "Фамилия Имя Отчество" in text or "фамилия, имя, отчество" in text.lower():
                replace_text_in_paragraph(paragraph, text.split()[0] if text.split() else text, STUDENT_FIO)
            
            # Заполняем группу
            if "группы" in text.lower() and "№" in text:
                if "МВА" not in text:
                    replace_text_in_paragraph(paragraph, text, text.replace("№", f"№ {STUDENT_GROUP}"))
            
            # Заполняем тему
            if "Тема практики" in text or "тема:" in text.lower():
                if "Разработка" not in text:
                    replace_text_in_paragraph(paragraph, text, f"{text}: {THEMA_PRACTICE}")
            
            # Заполняем организацию
            if "организация" in text.lower() or "место практики" in text.lower():
                if "VeritasAd" not in text:
                    replace_text_in_paragraph(paragraph, text, f"{text} - {ORGANIZATION}")
            
            # Заполняем даты
            if "с" in text and "по" in text and ("202" in text or "20" in text):
                if "2026" not in text:
                    new_text = text.replace(text.split("с")[-1].split("по")[-1].strip() if "с" in text and "по" in text else "", 
                                           f"{START_DATE.strftime('%d.%m.%Y')} по {END_DATE.strftime('%d.%m.%Y')}")
                    replace_text_in_paragraph(paragraph, text, new_text)
        
        # Сохраняем с новым именем
        output_path = os.path.join(DOCS_DIR, "4-Титульный лист отчета преддипломной практики Форма А_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Титульный лист сохранен: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении титульного листа: {e}")


def fill_diary():
    """Заполнение дневника практики (Приложение 5)"""
    print("\nЗаполнение дневника практики...")
    
    doc_path = os.path.join(DOCS_DIR, "2-Дневник преддипломной практики Приложение 5.docx")
    
    # План работ по дням
    daily_tasks = [
        ("09.02.2026", "Знакомство с проектом VeritasAd. Изучение архитектуры системы, документации и кодовой базы."),
        ("10.02.2026", "Настройка локального окружения для разработки. Установка Docker, Node.js, Python зависимостей."),
        ("11.02.2026", "Изучение API эндпоинтов системы. Анализ структуры базы данных и моделей."),
        ("12.02.2026", "Разработка новых функций для анализа видео. Интеграция ML-моделей Whisper и CLIP."),
        ("13.02.2026", "Оптимизация процесса обработки видео. Настройка кэширования результатов анализа."),
        ("16.02.2026", "Разработка системы тарифов и оплаты. Реализация pay-as-you-go модели."),
        ("17.02.2026", "Интеграция Telegram бота для взаимодействия с пользователями. Настройка авторизации."),
        ("18.02.2026", "Реализация системы аудита действий пользователей. Логирование событий безопасности."),
        ("19.02.2026", "Тестирование API эндпоинтов. Написание unit-тестов для критических функций."),
        ("20.02.2026", "Оптимизация производительности системы. Устранение уязвимостей безопасности."),
        ("21.02.2026", "Подготовка документации проекта. Обновление README и API документации."),
        ("22.02.2026", "Завершение работ по практике. Подготовка отчета и презентационных материалов."),
    ]
    
    try:
        doc = Document(doc_path)
        
        # Заполняем основную информацию
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Заполняем ФИО
            if "фамилия, имя, отчество" in text.lower():
                replace_text_in_paragraph(paragraph, text, STUDENT_FIO)
            
            # Заполняем группу
            if "группы" in text.lower() and "№" in text:
                if "МВА-122" not in text:
                    replace_text_in_paragraph(paragraph, text, text.replace("№", f"№ {STUDENT_GROUP}"))
        
        # Находим таблицу с дневником и заполняем
        if doc.tables:
            # Ищем таблицу с датами
            for table in doc.tables:
                # Проверяем, есть ли в таблице заголовки с датами
                if len(table.columns) >= 2:
                    first_row_text = table.rows[0].cells[0].text if table.rows else ""
                    if "Дата" in first_row_text or "№ п/п" in first_row_text:
                        # Заполняем строки таблицы
                        for i, (date, task) in enumerate(daily_tasks):
                            if i + 1 < len(table.rows):  # +1 потому что первая строка заголовок
                                row = table.rows[i + 1]
                                
                                # Заполняем ячейки
                                if len(row.cells) >= 2:
                                    row.cells[0].text = date  # Дата
                                    if len(row.cells) >= 2:
                                        row.cells[1].text = task  # Содержание работы
                                    if len(row.cells) >= 3:
                                        row.cells[2].text = ""    # Подпись (оставляем пустой)
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "2-Дневник преддипломной практики Приложение 5_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Дневник практики сохранен: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении дневника: {e}")


def fill_supervisor_conclusion():
    """Заполнение заключения руководителя (Форма Г)"""
    print("\nЗаполнение заключения руководителя...")
    
    doc_path = os.path.join(DOCS_DIR, "3-Заключение руководителя преддипломной Форма Г.docx")
    
    try:
        doc = Document(doc_path)
        
        # Заполняем основную информацию
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Заполняем ФИО студента
            if "обучающегося" in text.lower() or "студент" in text.lower():
                if STUDENT_FIO not in text:
                    # Находим место для вставки ФИО
                    if "фамилия" in text.lower():
                        replace_text_in_paragraph(paragraph, text, STUDENT_FIO)
            
            # Заполняем группу
            if "группы" in text.lower() and "№" in text:
                if STUDENT_GROUP not in text:
                    replace_text_in_paragraph(paragraph, text, text.replace("№", f"№ {STUDENT_GROUP}"))
            
            # Заполняем организацию
            if "организация" in text.lower() and "название" in text.lower():
                if ORGANIZATION not in text:
                    replace_text_in_paragraph(paragraph, text, f"{ORGANIZATION}")
            
            # Заполняем даты
            if "период" in text.lower() or ("с" in text and "по" in text):
                if "2026" not in text and ("202" in text or "20" in text):
                    new_text = f"с {START_DATE.strftime('%d.%m.%Y')} по {END_DATE.strftime('%d.%m.%Y')}"
                    replace_text_in_paragraph(paragraph, text, new_text)
        
        # Добавляем характеристику
        characteristic_text = (
            f"Студент {STUDENT_FIO} группы {STUDENT_GROUP} прошел преддипломную практику "
            f"в проекте {ORGANIZATION} в период с {START_DATE.strftime('%d.%m.%Y')} по {END_DATE.strftime('%d.%m.%Y')}. "
            f"За время практики студент продемонстрировал глубокие знания в области разработки веб-приложений, "
            f"работы с нейросетевыми моделями и базами данных. Были успешно реализованы ключевые функции системы "
            f"анализа видео VeritasAd, включая интеграцию ML-моделей (Whisper, CLIP), разработку системы тарифов "
            f"и pay-as-you-go модели, Telegram-интеграцию, систему аудита действий пользователей и обеспечения безопасности API."
        )
        
        # Находим параграф для характеристики и добавляем текст
        for paragraph in doc.paragraphs:
            if "Характеристика" in paragraph.text or "теоретических знаний" in paragraph.text:
                # Добавляем после этого параграфа
                if "Студент" not in paragraph.text:
                    p = paragraph.insert_paragraph_before(characteristic_text)
                    p.style = 'Normal'
                break
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "3-Заключение руководителя преддипломной Форма Г_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Заключение руководителя сохранено: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении заключения: {e}")


def fill_assignment():
    """Заполнение бланка задания на ВКР"""
    print("\nЗаполнение задания на ВКР...")
    
    doc_path = os.path.join(DOCS_DIR, "5-Бланк_задания.docx")
    
    try:
        doc = Document(doc_path)
        
        # Заполняем основную информацию
        for paragraph in doc.paragraphs:
            text = paragraph.text
            
            # Заполняем ФИО студента
            if "студент" in text.lower() or "фамилия" in text.lower():
                if STUDENT_FIO not in text and len(text.strip()) > 5:
                    # Ищем строку с ФИО
                    if any(char.isupper() for char in text):
                        replace_text_in_paragraph(paragraph, text, STUDENT_FIO)
            
            # Заполняем группу
            if "группа" in text.lower() or "группы" in text.lower():
                if STUDENT_GROUP not in text:
                    replace_text_in_paragraph(paragraph, text, f"группа {STUDENT_GROUP}")
            
            # Заполняем тему
            if "тема" in text.lower() and ("ВКР" in text or "диплом" in text.lower()):
                if THEMA_VKR not in text:
                    replace_text_in_paragraph(paragraph, text, f"Тема: {THEMA_VKR}")
        
        # Заполняем таблицы
        if doc.tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        # Заполняем ФИО в таблицах
                        if "ФИО" in text or "Фамилия" in text:
                            if STUDENT_FIO not in text:
                                cell.text = f"ФИО: {STUDENT_FIO}"
                        # Заполняем группу
                        if "группа" in text.lower():
                            if STUDENT_GROUP not in text:
                                cell.text = f"Группа: {STUDENT_GROUP}"
                        # Заполняем тему
                        if "тема" in text.lower():
                            if THEMA_VKR not in text:
                                cell.text = f"Тема: {THEMA_VKR}"
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "5-Бланк_задания_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Задание на ВКР сохранено: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении задания: {e}")


def main():
    """Основная функция"""
    print("=" * 60)
    print("Заполнение документов по преддипломной практике")
    print(f"Студент: {STUDENT_FIO}")
    print(f"Группа: {STUDENT_GROUP}")
    print(f"Период: {START_DATE.strftime('%d.%m.%Y')} - {END_DATE.strftime('%d.%m.%Y')}")
    print(f"Тема: {THEMA_PRACTICE}")
    print("=" * 60)
    
    fill_title_page()
    fill_diary()
    fill_supervisor_conclusion()
    fill_assignment()
    
    print("\n" + "=" * 60)
    print("Заполнение документов завершено!")
    print("Проверьте заполненные документы и при необходимости")
    print("отредактируйте их вручную в Microsoft Word.")
    print("=" * 60)


if __name__ == "__main__":
    main()
