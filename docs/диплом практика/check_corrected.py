"""
Скрипт для проверки 4 ИСПРАВЛЕННЫХ документов
"""

import os
from docx import Document

DOCS_DIR = r"C:\Users\dabin\Documents\Projects\VeritasAd\docs\диплом практика"

# 4 ИСПРАВЛЕННЫХ файла
files = [
    "4-Титульный лист отчета преддипломной практики Форма А_ИСПРАВЛЕН.docx",
    "2-Дневник преддипломной практики Приложение 5_ИСПРАВЛЕН.docx",
    "3-Заключение руководителя преддипломной Форма Г_ИСПРАВЛЕН.docx",
    "5-Бланк_задания_ИСПРАВЛЕН.docx",
]

def check_document(filename):
    """Проверка содержимого документа"""
    print(f"\n{'='*80}")
    print(f"ДОКУМЕНТ: {filename}")
    print('='*80)
    
    filepath = os.path.join(DOCS_DIR, filename)
    
    if not os.path.exists(filepath):
        print(f"[ERROR] Файл не найден: {filepath}")
        return None
    
    try:
        doc = Document(filepath)
        
        # Собираем весь текст из параграфов
        all_text = []
        print("\n--- Параграфы ---")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"{i+1}. {para.text}")
                all_text.append(para.text)
        
        # Собираем текст из таблиц
        if doc.tables:
            print(f"\n--- Таблицы ({len(doc.tables)}) ---")
            for t, table in enumerate(doc.tables):
                print(f"\nТаблица {t+1}:")
                for i, row in enumerate(table.rows):
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        print(f"  Строка {i+1}: {row_text}")
                        all_text.append(row_text)
        
        # Сохраняем полный текст для анализа
        full_text = '\n'.join(all_text)
        print(f"\n[OK] Документ прочитан. Всего символов: {len(full_text)}")
        
        return full_text
        
    except Exception as e:
        print(f"[ERROR] Ошибка при чтении документа: {e}")
        return None


def analyze_document(filename, full_text):
    """Анализ документа по критериям"""
    print(f"\n{'-'*80}")
    print(f"АНАЛИЗ: {filename}")
    print('-'*80)
    
    if full_text is None:
        print("[ERROR] Невозможно проанализировать - документ не прочитан")
        return
    
    # 1. Тема "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
    topic = "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
    topic_found = topic in full_text
    print(f"\n1. ТЕМА ПРАКТИКИ:")
    print(f"   '{topic}'")
    print(f"   Статус: {'[НАЙДЕНА]' if topic_found else '[НЕ НАЙДЕНА]'}")
    
    # 2. Заглушки "Петров В.Д." и "ООО Образец"
    petrov_found = "Петров В.Д." in full_text
    ooo_obrazets_found = "ООО Образец" in full_text
    print(f"\n2. ЗАГЛУШКИ (должны быть заменены):")
    print(f"   'Петров В.Д.' - {'[ЕСТЬ - ТРЕБУЕТ ЗАМЕНЫ]' if petrov_found else '[НЕТ - ОК]'}")
    print(f"   'ООО Образец' - {'[ЕСТЬ - ТРЕБУЕТ ЗАМЕНЫ]' if ooo_obrazets_found else '[НЕТ - ОК]'}")
    
    # 3. Пустые поля с подчеркиваниями
    import re
    empty_fields = re.findall(r'_{3,}', full_text)
    print(f"\n3. ПУСТЫЕ ПОЛЯ (подчеркивания):")
    print(f"   Количество: {len(empty_fields)}")
    if empty_fields:
        print(f"   Статус: '[ЕСТЬ ПУСТЫЕ ПОЛЯ]'")
    else:
        print(f"   Статус: '[НЕТ - ВСЕ ЗАПОЛНЕНО]'")
    
    # 4. Год 2025 (должен быть 2026)
    year_2025 = "2025" in full_text
    year_2026 = "2026" in full_text
    print(f"\n4. ГОД:")
    print(f"   2025 - {'[ЕСТЬ - ТРЕБУЕТ ЗАМЕНЫ на 2026]' if year_2025 else '[НЕТ - ОК]'}")
    print(f"   2026 - {'[ЕСТЬ - ОК]' if year_2026 else '[НЕТ]'}")
    
    print(f"\n{'-'*80}")


def main():
    """Проверка всех исправленных документов"""
    print("="*80)
    print("ПРОВЕРКА 4 ИСПРАВЛЕННЫХ ДОКУМЕНТОВ")
    print("="*80)
    
    all_texts = {}
    
    for filename in files:
        full_text = check_document(filename)
        all_texts[filename] = full_text
    
    print("\n\n" + "="*80)
    print("ИТОГОВЫЙ АНАЛИЗ ПО КРИТЕРИЯМ")
    print("="*80)
    
    for filename in files:
        analyze_document(filename, all_texts[filename])
    
    print("\n" + "="*80)
    print("ПРОВЕРКА ЗАВЕРШЕНА")
    print("="*80)


if __name__ == "__main__":
    main()
