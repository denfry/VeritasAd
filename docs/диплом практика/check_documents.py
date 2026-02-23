"""
Скрипт для проверки заполненных документов
"""

import os
from docx import Document

DOCS_DIR = r"C:\Users\dabin\Documents\Projects\VeritasAd\docs\диплом практика"

def check_document(filename):
    """Проверка содержимого документа"""
    print(f"\n{'='*60}")
    print(f"Документ: {filename}")
    print('='*60)
    
    filepath = os.path.join(DOCS_DIR, filename)
    
    if not os.path.exists(filepath):
        print(f"[ERROR] Файл не найден: {filepath}")
        return
    
    try:
        doc = Document(filepath)
        
        # Выводим параграфы
        print("\n--- Параграфы ---")
        for i, para in enumerate(doc.paragraphs[:20]):  # Первые 20 параграфов
            if para.text.strip():
                print(f"{i+1}. {para.text[:200]}")  # Ограничиваем длину
        
        # Выводим таблицы
        if doc.tables:
            print(f"\n--- Таблицы ({len(doc.tables)}) ---")
            for t, table in enumerate(doc.tables):
                print(f"\nТаблица {t+1} ({len(table.rows)} строк, {len(table.columns)} колонок):")
                for i, row in enumerate(table.rows[:5]):  # Первые 5 строк
                    cells_text = [cell.text[:50] for cell in row.cells]
                    print(f"  Строка {i+1}: {cells_text}")
        
        print(f"\n[OK] Документ проверен: {filename}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при чтении документа: {e}")


def main():
    """Проверка всех заполненных документов"""
    files = [
        "4-Титульный лист отчета преддипломной практики Форма А_ЗАПОЛНЕН.docx",
        "2-Дневник преддипломной практики Приложение 5_ЗАПОЛНЕН.docx",
        "3-Заключение руководителя преддипломной Форма Г_ЗАПОЛНЕН.docx",
        "5-Бланк_задания_ЗАПОЛНЕН.docx",
    ]
    
    print("="*60)
    print("ПРОВЕРКА ЗАПОЛНЕННЫХ ДОКУМЕНТОВ")
    print("="*60)
    
    for filename in files:
        check_document(filename)
    
    print("\n" + "="*60)
    print("ПРОВЕРКА ЗАВЕРШЕНА")
    print("="*60)


if __name__ == "__main__":
    main()
