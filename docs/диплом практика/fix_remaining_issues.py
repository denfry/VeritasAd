"""
Дополнительный скрипт для исправления оставшихся проблем в документах
Особое внимание: Форма Г - замена "Петров В.Д." и года 2025
"""

import os
import sys
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from datetime import datetime

STUDENT_FIO = "Юрков Данила Андреевич"
THEMA_PRACTICE = "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
ORGANIZATION = "VeritasAd"
START_DATE = datetime(2026, 2, 9)
END_DATE = datetime(2026, 2, 22)
SUPERVISOR = "Беспалов М.Е., к.т.н., доц."

DOCS_DIR = r"C:\Users\dabin\Documents\Projects\VeritasAd\docs\диплом практика"


def fix_form_g():
    """Исправление Формы Г - замена Петров В.Д. и года 2025"""
    print("Исправление Формы Г...")
    
    # Читаем ИСПРАВЛЕННЫЙ файл
    doc_path = os.path.join(DOCS_DIR, "3-Заключение руководителя преддипломной Форма Г_ИСПРАВЛЕН.docx")
    
    try:
        doc = Document(doc_path)
        
        changes_made = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text
            original_text = text
            
            # 1. Замена "Петров В.Д." и "Петрова В.Д."
            if "Петров В.Д." in text or "Петрова В.Д." in text:
                text = text.replace("Петров В.Д.", STUDENT_FIO)
                text = text.replace("Петрова В.Д.", STUDENT_FIO)
                changes_made += 1
                print(f"  [ИСПРАВЛЕНО] Петров В.Д. -> {STUDENT_FIO}")
            
            # 2. Исправление года 2025 на 2026 в датах
            if "2025 г." in text or "2025г." in text or "2025 г" in text:
                text = text.replace("2025 г.", "2026 г.")
                text = text.replace("2025г.", "2026 г.")
                text = text.replace("2025 г", "2026 г")
                changes_made += 1
                print(f"  [ИСПРАВЛЕНО] 2025 -> 2026")
            
            # 3. Исправление периода практики (январь 2025 -> февраль 2026)
            if "13» января 2025" in text or "09» февраля 2025" in text:
                text = text.replace("13» января 2025 г.", f"09» февраля 2026 г.")
                text = text.replace("09» февраля 2025 г.", f"22» февраля 2026 г.")
                changes_made += 1
                print(f"  [ИСПРАВЛЕНО] Период практики исправлен")
            
            # 4. Добавление темы практики если отсутствует
            if "Тема" in text and ("ВКР" in text or "практики" in text) and THEMA_PRACTICE not in text:
                text = f"Тема: {THEMA_PRACTICE}"
                changes_made += 1
                print(f"  [ДОБАВЛЕНО] Тема практики")
            
            # Применяем изменения если были
            if text != original_text:
                runs = paragraph.runs
                if runs:
                    first_run = runs[0]
                    for run in runs:
                        run.text = ""
                    runs[0].text = text
        
        # Сохраняем в НОВЫЙ файл с пометкой ФИНАЛЬНЫЙ
        output_path = os.path.join(DOCS_DIR, "3-Заключение руководителя преддипломной Форма Г_ФИНАЛЬНЫЙ.docx")
        doc.save(output_path)
        print(f"\n[OK] Форма Г сохранена: {output_path}")
        print(f"Всего исправлений: {changes_made}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при исправлении Формы Г: {e}")


def fix_form_a():
    """Исправление Формы А - добавление темы практики"""
    print("\nИсправление Формы А...")
    
    doc_path = os.path.join(DOCS_DIR, "4-Титульный лист отчета преддипломной практики Форма А_ИСПРАВЛЕН.docx")
    
    try:
        doc = Document(doc_path)
        
        changes_made = 0
        
        # Ищем место для добавления темы
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            
            # Добавляем тему после заголовка "ОТЧЕТ"
            if "ОТЧЕТ" in text and "преддипломной практике" in text:
                # Вставляем тему после этого параграфа
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if THEMA_PRACTICE not in next_para.text:
                        # Вставляем новый параграф с темой
                        new_para = doc.add_paragraph(THEMA_PRACTICE, next_para)
                        new_para.style = 'Normal'
                        changes_made += 1
                        print(f"  [ДОБАВЛЕНО] Тема практики")
                        break
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "4-Титульный лист отчета преддипломной практики Форма А_ФИНАЛЬНЫЙ.docx")
        doc.save(output_path)
        print(f"[OK] Форма А сохранена: {output_path}")
        print(f"Всего исправлений: {changes_made}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при исправлении Формы А: {e}")


def main():
    print("=" * 70)
    print("ДОПОЛНИТЕЛЬНОЕ ИСПРАВЛЕНИЕ ДОКУМЕНТОВ")
    print("=" * 70)
    print()
    
    fix_form_g()
    fix_form_a()
    
    print()
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ЗАВЕРШЕНО!")
    print("=" * 70)
    print()
    print("ФИНАЛЬНЫЕ файлы:")
    print("  1. 3-Заключение руководителя преддипломной Форма Г_ФИНАЛЬНЫЙ.docx")
    print("  2. 4-Титульный лист отчета преддипломной практики Форма А_ФИНАЛЬНЫЙ.docx")
    print()


if __name__ == "__main__":
    main()
