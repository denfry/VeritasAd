"""
Скрипт для заполнения документов по преддипломной практике
Проект: VeritasAd - AI платформа детекции скрытой рекламы
Студент: Юрков Данила Андреевич, группа МВА-122
Период практики: 09.02.2026 - 22.02.2026
"""

import os
import sys
# Устанавливаем UTF-8 кодировку для вывода в Windows
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta

# Константы
STUDENT_FIO = "Юрков Данила Андреевич"
STUDENT_GROUP = "МВА-122"
THEMA_PRACTICE = "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
THEMA_VKR = "Разработка подсистемы нейросетевого анализа артефактов рекламного видео"
ORGANIZATION = "VeritasAd"
START_DATE = datetime(2026, 2, 9)
END_DATE = datetime(2026, 2, 22)
SUPERVISOR = "Беспалов М.Е., к.т.н., доц."
HEAD_DEPARTMENT = "Травкин Е.И., к.п.н."

DOCS_DIR = r"C:\Users\dabin\Documents\Projects\VeritasAd\docs\диплом практика"


def fill_title_page():
    """Заполнение титульного листа отчета (Форма А)"""
    print("Заполнение титульного листа отчета...")
    
    doc_path = os.path.join(DOCS_DIR, "4-Титульный лист отчета преддипломной практики Форма А.docx")
    
    try:
        doc = Document(doc_path)
        
        # Заменяем текст в документе
        replacements = {
            "Юрков Данила Андреевич": STUDENT_FIO,
            "МВА-122": STUDENT_GROUP,
            "Разработка подсистемы нейросетевого анализа артефактов рекламного видео": THEMA_PRACTICE,
            "VeritasAd": ORGANIZATION,
            "09.02.2026": START_DATE.strftime("%d.%m.%Y"),
            "22.02.2026": END_DATE.strftime("%d.%m.%Y"),
        }
        
        # Проходим по всем параграфам и заменяем текст
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    # Сохраняем форматирование первого запуска
                    runs = paragraph.runs
                    if runs:
                        # Очищаем параграф и добавляем новый текст с тем же форматированием
                        original_run = runs[0]
                        for run in runs:
                            run.text = ""
                        runs[0].text = paragraph.text.replace(old_text, new_text)
        
        # Сохраняем с новым именем
        output_path = os.path.join(DOCS_DIR, "4-Титульный лист отчета преддипломной практики Форма А_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Титульный лист сохранен: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении титульного листа: {e}")


def fill_diary():
    """Заполнение дневника практики (Приложение 5)"""
    print("\nЗаполнение дневника практики...")
    
    doc_path = os.path.join(DOCS_DIR, "2-Дневник преддипломной практики Приложение 5.docx")
    
    # План работ по дням
    daily_tasks = [
        ("09.02.2026", "Знакомство с проектом VeritasAd. Изучение архитектуры системы, документации и кодовой базы."),
        ("10.02.2026", "Настройка локального окружения для разработки. Установка Docker, Node.js, Python зависимостей."),
        ("11.02.2026", "Изучение API эндпоинтов системы. Анализ структуры базы данных и моделей."),
        ("12.02.2026", "Разработка новых функций для анализа видео. Интеграция ML-моделей Whisper и CLIP."),
        ("13.02.2026", "Оптимизация процесса обработки видео. Настройка кэширования результатов анализа."),
        ("16.02.2026", "Разработка системы тарифов и оплаты. Реализация pay-as-you-go модели."),
        ("17.02.2026", "Интеграция Telegram бота для взаимодействия с пользователями. Настройка авторизации."),
        ("18.02.2026", "Реализация системы аудита действий пользователей. Логирование событий безопасности."),
        ("19.02.2026", "Тестирование API эндпоинтов. Написание unit-тестов для критических функций."),
        ("20.02.2026", "Оптимизация производительности системы. Устранение уязвимостей безопасности."),
        ("21.02.2026", "Подготовка документации проекта. Обновление README и API документации."),
        ("22.02.2026", "Завершение работ по практике. Подготовка отчета и презентационных материалов."),
    ]
    
    try:
        doc = Document(doc_path)
        
        # Находим таблицу с дневником
        if doc.tables:
            table = doc.tables[0]  # Предполагаем, что первая таблица - это дневник
            
            # Заполняем строки таблицы (начиная со второй строки, т.к. первая - заголовок)
            for i, (date, task) in enumerate(daily_tasks):
                if i + 1 < len(table.rows):  # +1 потому что первая строка заголовок
                    row = table.rows[i + 1]
                    
                    # Заполняем ячейки
                    if len(row.cells) >= 3:
                        row.cells[0].text = date  # Дата
                        row.cells[1].text = task  # Содержание работы
                        row.cells[2].text = ""    # Подпись (оставляем пустой)
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "2-Дневник преддипломной практики Приложение 5_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Дневник практики сохранен: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении дневника: {e}")


def fill_supervisor_conclusion():
    """Заполнение заключения руководителя (Форма Г)"""
    print("\nЗаполнение заключения руководителя...")
    
    doc_path = os.path.join(DOCS_DIR, "3-Заключение руководителя преддипломной Форма Г.docx")
    
    try:
        doc = Document(doc_path)
        
        # Заменяем текст
        replacements = {
            "Юрков Данила Андреевич": STUDENT_FIO,
            "МВА-122": STUDENT_GROUP,
        }
        
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    runs = paragraph.runs
                    if runs:
                        for run in runs:
                            run.text = ""
                        runs[0].text = paragraph.text.replace(old_text, new_text)
        
        # Добавляем характеристику если есть место
        characteristic = (
            f"Студент {STUDENT_FIO} группы {STUDENT_GROUP} прошел преддипломную практику "
            f"в проекте {ORGANIZATION} в период с {START_DATE.strftime('%d.%m.%Y')} по {END_DATE.strftime('%d.%m.%Y')}. "
            f"За время практики студент продемонстрировал глубокие знания в области разработки веб-приложений, "
            f"работы с нейросетевыми моделями и базами данных. Были успешно реализованы ключевые функции системы "
            f"анализа видео, включая интеграцию ML-моделей, систему тарифов и Telegram-интеграцию."
        )
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "3-Заключение руководителя преддипломной Форма Г_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Заключение руководителя сохранено: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении заключения: {e}")


def fill_assignment():
    """Заполнение бланка задания на ВКР"""
    print("\nЗаполнение задания на ВКР...")
    
    doc_path = os.path.join(DOCS_DIR, "5-Бланк_задания.docx")
    
    try:
        doc = Document(doc_path)
        
        # Заменяем текст
        replacements = {
            "Юрков Данила Андреевич": STUDENT_FIO,
            "МВА-122": STUDENT_GROUP,
            "Разработка подсистемы нейросетевого анализа артефактов рекламного видео": THEMA_VKR,
        }
        
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    runs = paragraph.runs
                    if runs:
                        for run in runs:
                            run.text = ""
                        runs[0].text = paragraph.text.replace(old_text, new_text)
        
        # Сохраняем
        output_path = os.path.join(DOCS_DIR, "5-Бланк_задания_ЗАПОЛНЕН.docx")
        doc.save(output_path)
        print(f"[OK] Задание на ВКР сохранено: {output_path}")
        
    except Exception as e:
        print(f"[ERROR] Ошибка при заполнении задания: {e}")


def main():
    """Основная функция"""
    print("=" * 60)
    print("Заполнение документов по преддипломной практике")
    print(f"Студент: {STUDENT_FIO}")
    print(f"Группа: {STUDENT_GROUP}")
    print(f"Период: {START_DATE.strftime('%d.%m.%Y')} - {END_DATE.strftime('%d.%m.%Y')}")
    print("=" * 60)
    
    fill_title_page()
    fill_diary()
    fill_supervisor_conclusion()
    fill_assignment()
    
    print("\n" + "=" * 60)
    print("Заполнение документов завершено!")
    print("=" * 60)


if __name__ == "__main__":
    main()
